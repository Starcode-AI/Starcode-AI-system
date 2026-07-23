import hashlib
import io
import json
import mimetypes
import re
import socket
import struct
import tarfile
import zipfile
from dataclasses import asdict, dataclass
from pathlib import Path, PurePosixPath

from pypdf import PdfReader

from .code_review import review_code
from .safety import sanitize_untrusted_text


TEXT_EXTENSIONS = {
    ".txt", ".md", ".json", ".csv", ".xml", ".yaml", ".yml", ".html", ".css",
    ".js", ".ts", ".py", ".php", ".lua", ".sql", ".log", ".java", ".cs", ".c",
    ".h", ".cpp", ".hpp", ".rs", ".go", ".ps1", ".sh",
}
ARCHIVE_EXTENSIONS = {".zip", ".tar", ".gz", ".tgz", ".7z"}
DANGEROUS_EXTENSIONS = {".exe", ".dll", ".msi", ".com", ".scr", ".bat", ".cmd", ".vbs", ".jscript", ".jar"}


@dataclass
class ArchiveReport:
    safe: bool
    file_count: int
    expanded_bytes: int
    compressed_bytes: int
    ratio: float
    issues: list[str]
    entries: list[str]


def scan_with_clamav(path: Path, host: str, port: int = 3310) -> dict:
    if not host:
        return {"enabled": False, "status": "not_configured"}
    try:
        with socket.create_connection((host, port), timeout=10) as connection:
            connection.sendall(b"zINSTREAM\0")
            with path.open("rb") as handle:
                while chunk := handle.read(1024 * 1024):
                    connection.sendall(struct.pack(">I", len(chunk)))
                    connection.sendall(chunk)
            connection.sendall(struct.pack(">I", 0))
            response = connection.recv(4096).rstrip(b"\0").decode("utf-8", errors="replace")
        infected = response.endswith(" FOUND")
        return {"enabled": True, "status": "infected" if infected else "clean", "response": response[:500]}
    except OSError as exc:
        return {"enabled": True, "status": "unavailable", "error": str(exc)[:200]}


def sanitize_filename(name: str) -> str:
    base = Path(name.replace("\\", "/")).name
    base = re.sub(r"[^A-Za-z0-9ÄÖÜäöüß._ -]", "_", base).strip(" .")
    return base[:180] or "upload.bin"


def has_suspicious_double_extension(name: str) -> bool:
    suffixes = [item.lower() for item in Path(name).suffixes]
    return len(suffixes) >= 2 and any(item in DANGEROUS_EXTENSIONS for item in suffixes)


def safe_archive_path(name: str) -> bool:
    path = PurePosixPath(name.replace("\\", "/"))
    return not path.is_absolute() and ".." not in path.parts and not re.match(r"^[A-Za-z]:", name)


def inspect_zip(path: Path, max_files: int, max_expanded: int) -> ArchiveReport:
    issues: list[str] = []
    entries: list[str] = []
    expanded = 0
    compressed = 0
    try:
        with zipfile.ZipFile(path) as archive:
            infos = archive.infolist()
            if len(infos) > max_files:
                issues.append("Archive contains too many files")
            for info in infos[:max_files + 1]:
                if not safe_archive_path(info.filename):
                    issues.append(f"Unsafe path: {info.filename[:120]}")
                mode = (info.external_attr >> 16) & 0o170000
                if mode == 0o120000:
                    issues.append(f"Symbolic link blocked: {info.filename[:120]}")
                expanded += info.file_size
                compressed += info.compress_size
                entries.append(info.filename[:240])
                if expanded > max_expanded:
                    issues.append("Expanded archive exceeds the configured limit")
                    break
    except (zipfile.BadZipFile, OSError) as exc:
        issues.append(f"Damaged ZIP archive: {exc}")
    ratio = expanded / max(compressed, 1)
    if ratio > 200 and expanded > 50 * 1024 * 1024:
        issues.append("Suspicious compression ratio; possible ZIP bomb")
    return ArchiveReport(not issues, len(entries), expanded, compressed, ratio, issues, entries[:500])


def inspect_tar(path: Path, max_files: int, max_expanded: int) -> ArchiveReport:
    issues: list[str] = []
    entries: list[str] = []
    expanded = 0
    try:
        with tarfile.open(path, "r:*") as archive:
            for index, info in enumerate(archive):
                if index >= max_files:
                    issues.append("Archive contains too many files")
                    break
                if not safe_archive_path(info.name):
                    issues.append(f"Unsafe path: {info.name[:120]}")
                if info.issym() or info.islnk() or info.isdev():
                    issues.append(f"Links and device entries are blocked: {info.name[:120]}")
                expanded += max(info.size, 0)
                entries.append(info.name[:240])
                if expanded > max_expanded:
                    issues.append("Expanded archive exceeds the configured limit")
                    break
    except (tarfile.TarError, OSError) as exc:
        issues.append(f"Damaged TAR archive: {exc}")
    compressed = path.stat().st_size if path.exists() else 0
    ratio = expanded / max(compressed, 1)
    if ratio > 200 and expanded > 50 * 1024 * 1024:
        issues.append("Suspicious compression ratio; possible archive bomb")
    return ArchiveReport(not issues, len(entries), expanded, compressed, ratio, issues, entries[:500])


def inspect_7z(path: Path, max_files: int, max_expanded: int) -> ArchiveReport:
    try:
        import py7zr
    except ImportError:
        return ArchiveReport(False, 0, 0, path.stat().st_size, 0.0, ["7Z support is not installed"], [])
    issues: list[str] = []
    entries: list[str] = []
    expanded = 0
    try:
        with py7zr.SevenZipFile(path, mode="r") as archive:
            infos = archive.list()
            if len(infos) > max_files:
                issues.append("Archive contains too many files")
            for info in infos[:max_files + 1]:
                name = str(info.filename)
                if not safe_archive_path(name):
                    issues.append(f"Unsafe path: {name[:120]}")
                expanded += int(getattr(info, "uncompressed", 0) or 0)
                entries.append(name[:240])
                if expanded > max_expanded:
                    issues.append("Expanded archive exceeds the configured limit")
                    break
    except Exception as exc:
        issues.append(f"Damaged or encrypted 7Z archive: {str(exc)[:160]}")
    compressed = path.stat().st_size
    ratio = expanded / max(compressed, 1)
    if ratio > 200 and expanded > 50 * 1024 * 1024:
        issues.append("Suspicious compression ratio; possible archive bomb")
    return ArchiveReport(not issues, len(entries), expanded, compressed, ratio, issues, entries[:500])


def analyze_pdf(path: Path) -> dict:
    try:
        reader = PdfReader(str(path), strict=False)
    except Exception as exc:
        return {"type": "pdf", "error": f"PDF could not be read: {str(exc)[:200]}"}
    if reader.is_encrypted:
        try:
            if reader.decrypt("") == 0:
                return {"type": "pdf", "encrypted": True, "error": "PDF is encrypted"}
        except Exception:
            return {"type": "pdf", "encrypted": True, "error": "PDF is encrypted"}
    pages: list[dict] = []
    for number, page in enumerate(reader.pages[:1000], start=1):
        try:
            text = page.extract_text() or ""
        except Exception as exc:
            text = f"[Page could not be extracted: {str(exc)[:100]}]"
        clean, injection = sanitize_untrusted_text(text, 100_000)
        pages.append({"page": number, "text": clean, "injection_detected": injection})
    metadata = {str(k): str(v)[:1000] for k, v in (reader.metadata or {}).items()}
    return {"type": "pdf", "encrypted": False, "page_count": len(reader.pages), "metadata": metadata, "pages": pages}


def analyze_docx(path: Path) -> dict:
    try:
        from docx import Document

        document = Document(path)
        parts = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
        tables = []
        for table in document.tables[:100]:
            tables.append([[cell.text[:5000] for cell in row.cells] for row in table.rows[:1000]])
        clean, injection = sanitize_untrusted_text("\n".join(parts), 2_000_000)
        return {"type": "docx", "text": clean, "tables": tables, "prompt_injection_detected": injection}
    except Exception as exc:
        return {"type": "docx", "error": f"DOCX could not be read: {str(exc)[:200]}"}


def analyze_xlsx(path: Path) -> dict:
    try:
        from openpyxl import load_workbook

        workbook = load_workbook(path, read_only=True, data_only=True)
        sheets = []
        text_parts = []
        for sheet in workbook.worksheets[:50]:
            rows = []
            for index, row in enumerate(sheet.iter_rows(values_only=True)):
                if index >= 5000:
                    break
                values = ["" if value is None else str(value)[:5000] for value in row[:200]]
                rows.append(values)
                text_parts.append(" | ".join(values))
            sheets.append({"title": sheet.title, "rows": rows})
        clean, injection = sanitize_untrusted_text("\n".join(text_parts), 2_000_000)
        return {"type": "xlsx", "sheets": sheets, "text": clean, "prompt_injection_detected": injection}
    except Exception as exc:
        return {"type": "xlsx", "error": f"XLSX could not be read: {str(exc)[:200]}"}


def analyze_file(path: Path, original_name: str, max_files: int, max_expanded: int) -> dict:
    suffix = path.suffix.lower()
    guessed_mime = mimetypes.guess_type(original_name)[0] or "application/octet-stream"
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(chunk)
    result: dict = {
        "filename": original_name,
        "safe_filename": sanitize_filename(original_name),
        "size": path.stat().st_size,
        "sha256": digest.hexdigest(),
        "mime": guessed_mime,
        "double_extension_warning": has_suspicious_double_extension(original_name),
        "executable": suffix in DANGEROUS_EXTENSIONS,
    }
    if suffix == ".pdf":
        result["analysis"] = analyze_pdf(path)
    elif suffix in {".docx", ".xlsx"}:
        archive_check = inspect_zip(path, max_files, max_expanded)
        if not archive_check.safe:
            result["analysis"] = {
                "type": suffix.removeprefix("."),
                "error": "Office document failed its archive safety check",
                "archive_check": asdict(archive_check),
            }
        elif suffix == ".docx":
            result["analysis"] = {**analyze_docx(path), "archive_check": asdict(archive_check)}
        else:
            result["analysis"] = {**analyze_xlsx(path), "archive_check": asdict(archive_check)}
    elif suffix == ".zip":
        result["analysis"] = {"type": "archive", **asdict(inspect_zip(path, max_files, max_expanded))}
    elif suffix in {".tar", ".gz", ".tgz"} or original_name.lower().endswith(".tar.gz"):
        result["analysis"] = {"type": "archive", **asdict(inspect_tar(path, max_files, max_expanded))}
    elif suffix == ".7z":
        result["analysis"] = {"type": "archive", **asdict(inspect_7z(path, max_files, max_expanded))}
    elif suffix in TEXT_EXTENSIONS:
        raw = path.read_bytes()
        text = raw.decode("utf-8", errors="replace")[:2_000_000]
        clean, injection = sanitize_untrusted_text(text, 2_000_000)
        result["analysis"] = {
            "type": "text",
            "text": clean,
            "line_count": clean.count("\n") + 1,
            "prompt_injection_detected": injection,
            "code_review": review_code(clean, original_name) if suffix not in {".txt", ".md", ".log", ".csv"} else None,
        }
    else:
        result["analysis"] = {"type": "binary", "message": "Stored for inspection but never executed"}
    return result
